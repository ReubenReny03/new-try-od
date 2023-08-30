
from flask import Flask, render_template, request, send_file
import pandas as pd
from docx import Document
import io
import base64
import os

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)

def change_names_new_doc(name,date,session,template_document_p):
    variables = {
        "${event_name}": name,
        "${date}": date,
        "${session}": session,
    }
    template_document = Document(template_document_p)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)
    #! what the heking code is this
    # doc_stream = io.BytesIO()
    # template_document.save(doc_stream)
    # doc_stream.seek(0)
    # return doc_stream

    # return send_file(
    #     doc_stream,
    #     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    #     as_attachment=True,
    #     download_name='my_document.docx'
    # )
    return template_document

def od_list(reg_f,event_f,doc):
    reg_file = pd.read_csv(reg_f)
    event_file = pd.read_csv(event_f)

    given_list = list(event_file["id"])

    new_list = pd.DataFrame()
    
    for x in given_list:
        a = reg_file.loc[reg_file["Register Number"].str.upper() == x.upper()]
        
        if not a.empty:  # Check if there are any records for the current x
            name = list(a["Full Name (in all capital letters)"])
            department = list(a["Department of Study"])
            data = {"name": f"{name[0]}", "Registration Number": x, "Department": f"{department[0]}"}
            new_list = new_list._append(data, ignore_index=True)
        else:
            print(f"No records found for Register Number: {x}")
    print(new_list)
    table = doc.tables[1]
    count = 1
    for index, row in new_list.iterrows():
        s_no = count
        count +=1
        name = row["name"]
        registration_number = row["Registration Number"]
        department = row["Department"]
        # Assuming you have three columns in your table

        table.add_row().cells[0].text = str(s_no)
        table.rows[-1].cells[1].text = name
        table.rows[-1].cells[2].text = str(registration_number)
        table.rows[-1].cells[3].text = department

    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream
    # a = new_list.to_csv("output.csv", index=True)

app = Flask(__name__)
#! Remember attendect --> id
#! Remember registry --> Register Number, Full Name (in all capital letters), Department of Study, 

@app.route('/',methods=['GET','POST'])
def root():
    if request.method == "POST":
        # getting input with name = fname in HTML form
        event_name = request.form.get("name")
        # getting input with name = lname in HTML form
        date = request.form.get("date")
        session = request.form.get("session")
        attendece_list = request.files["a_list"]
        reg_list = request.files["r_list"]
        doc_demo = request.files["t_doc"]
        new_doc = change_names_new_doc(event_name,date,session,doc_demo)
        doc_stream = od_list(reg_list,attendece_list,new_doc)
        return send_file(
        doc_stream,
        as_attachment=True,
        download_name="od_list.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
        #return "Your data is sent to the host"
    return render_template("main.html")


if __name__ == '__main__':
    app.run(port=5001, debug=True)